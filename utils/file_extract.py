"""上传文件 → 文本 的统一提取层。

⚠️ 这个模块存在的理由是【不要有第二份实现】。原本 page 2(新建项目)和
page 3(请旨补充资料)各写了一套上传处理:page 2 认真解析 pdf/docx,page 3
却对一切非图片文件直接 `data.decode("utf-8")` —— 传一份 Word 进去,brief
里就多一大坨替换字符。两边的体积上限也是各写各的常量,改一处永远漏另一处。

上一轮评审刚记过同一个教训(采样手写了一个平行的重试分类器,把已修的 bug
重新引入了一遍)。**复制行为比复制代码更容易漏掉补丁**,所以这里合并成一份。

入参是任何带 `.name` / `.read()`(可选 `.size`)的对象 —— Streamlit 的
UploadedFile 正好是这个形状。
"""

import io

# Hard caps enforced server-side (the `type=[...]` arg on st.file_uploader is
# just a UI hint, the client can still post anything). Images go into the
# prompt as base64, which inflates by ~33%; 2 MB decoded → ~2.7 MB encoded,
# which is already a lot of tokens. Text-like files are bounded to 1 MB so a
# runaway upload can't blow the free-text column / LLM context.
_MAX_IMAGE_BYTES = 2 * 1024 * 1024
_MAX_TEXT_BYTES = 1 * 1024 * 1024
# ⚠️ 容器类文件(docx / pdf)的体积上限**必须和文本上限分开量**。
#
# 这两种格式是压缩包/复合文档,里面装的绝大部分是图片字节,和它能吐出多少
# 文字几乎无关。拿 1MB 的文本上限去卡整个容器,后果是【一份图文并茂的 Word
# 会因为配图太大被整份拒收,哪怕正文只有几 KB】—— 真实案例:4.01MB 的样稿
# 文档,16 张配图占 4.4MB,正文 7,827 字(约 8KB),被 1MB 上限整份扔掉,
# 用户只在清单里看到一行 33 字的报错。上限本来想拦的是"文本撑爆上下文",
# 结果拦掉的是"文档里有图"。
#
# 所以:容器给一个宽松的读入上限(挡住真正离谱的上传),**提取出来的文本
# 再单独按 _MAX_TEXT_BYTES 截断** —— 这才是上限原本想管的那个量。
_MAX_CONTAINER_BYTES = 30 * 1024 * 1024
_TEXT_EXTS = (".txt", ".md", ".json")
_IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".gif", ".webp")


def _read_with_cap(uploaded_file, max_bytes: int) -> bytes | None:
    """Read uploaded_file and return bytes, or None if it exceeds max_bytes.
    Uses .size when available (Streamlit UploadedFile exposes it) to avoid
    materializing oversized content in memory."""
    size = getattr(uploaded_file, "size", None)
    if size is not None and size > max_bytes:
        return None
    data = uploaded_file.read()
    if len(data) > max_bytes:
        return None
    return data


def _cap_extracted_text(text: str, filename: str) -> str:
    """把**提取出来的文本**截到 _MAX_TEXT_BYTES,并明说截断了。

    截断而不是整份丢弃:文本超限时前面那部分仍然有价值,而整份拒收等于
    用户什么都拿不到(还只看得到一行报错)。
    """
    encoded = text.encode("utf-8")
    if len(encoded) <= _MAX_TEXT_BYTES:
        return text
    cut = encoded[:_MAX_TEXT_BYTES].decode("utf-8", errors="ignore")
    return (
        cut
        + f"\n\n[⚠️ {filename} 提取出的文本超过 {_MAX_TEXT_BYTES // 1024}KB，"
          f"已截断；后半部分未进入 brief]"
    )


def _count_docx_images(data: bytes) -> int:
    """数 docx 里内嵌了多少张图(zip 里的 word/media/ 条目)。

    不用 python-docx 的 inline_shapes:那个只数内联图,浮动图(环绕排版)
    数不到,而中文文档里浮动图很常见。
    """
    try:
        import zipfile
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            return sum(
                1 for n in z.namelist()
                if n.startswith("word/media/") and not n.endswith("/")
            )
    except Exception:
        return 0


def _docx_to_text(data: bytes, filename: str) -> str:
    """按文档顺序提取 docx 的段落 + 表格，并报告未提取的内嵌图。

    ⚠️ 不能只取 `doc.paragraphs` —— 那个**不含表格单元格**。而产品资料类
    Word 的硬事实(成分 / 规格 / 价格带 / 竞品对比)基本全在表格里,只取
    段落等于把最该进 brief 的部分丢掉,且丢得静默。
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(data))
    para_by_el = {id(p._p): p for p in doc.paragraphs}
    tbl_by_el = {id(t._tbl): t for t in doc.tables}

    out: list[str] = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = para_by_el.get(id(child))
            if p is None:
                continue
            txt = p.text.strip()
            if not txt:
                continue
            style = (p.style.name or "") if p.style is not None else ""
            if style.startswith("Heading"):
                lvl = "".join(c for c in style if c.isdigit()) or "1"
                out.append("\n" + "#" * min(int(lvl), 6) + " " + txt)
            else:
                out.append(txt)
        elif child.tag == qn("w:tbl"):
            t = tbl_by_el.get(id(child))
            if t is None:
                continue
            rows = []
            for r in t.rows:
                cells = [c.text.strip().replace("\n", " ") for c in r.cells]
                if any(cells):
                    rows.append(cells)
            if not rows:
                # 纯排版用的空表格(中文 Word 里很常见),跳过不占位
                continue
            out.append("")
            out.append("| " + " | ".join(rows[0]) + " |")
            out.append("|" + "---|" * len(rows[0]))
            for r in rows[1:]:
                out.append("| " + " | ".join(r) + " |")
            out.append("")

    text = "\n".join(out)

    # 内嵌图不做转写(一份文档十几张图 = 几分钟的 Vision 调用,而 Word 里的
    # 图多半是配图不承载信息)。但**必须明说没提取**,否则用户以为图进去了。
    n_img = _count_docx_images(data)
    if n_img:
        text += (
            f"\n\n[⚠️ 本文档内嵌 {n_img} 张图片，未被提取。"
            f"如果图里有信息（截图 / 数据图 / 竞品对比），请把它们导出后"
            f"作为单独的图片文件上传——那条路径会走 Vision 转写。]"
        )
    return text


def extract_file_content(uploaded_file) -> str:
    """Extract text content from an uploaded file. Enforces size caps."""
    name = uploaded_file.name.lower()

    if name.endswith(_TEXT_EXTS):
        data = _read_with_cap(uploaded_file, _MAX_TEXT_BYTES)
        if data is None:
            return f"[文件过大跳过: {uploaded_file.name} > {_MAX_TEXT_BYTES // 1024}KB]"
        return data.decode("utf-8", errors="replace")

    if name.endswith(".pdf"):
        data = _read_with_cap(uploaded_file, _MAX_CONTAINER_BYTES)
        if data is None:
            return (
                f"[PDF 过大跳过: {uploaded_file.name} > "
                f"{_MAX_CONTAINER_BYTES // 1024 // 1024}MB]"
            )
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(data))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            return f"[PDF 解析失败: {e}]"
        if not text.strip():
            # 扫描版 PDF 没有文字层,PyPDF2 会安静地返回空串。不说明的话
            # 用户只会看到"已提取 0 字",完全不知道该怎么办。
            return (
                f"[PDF 无文字层: {uploaded_file.name} —— 多半是扫描件或纯图排版，"
                f"PyPDF2 提取不到任何文字。请把页面导出成图片单独上传，"
                f"那条路径会走 Vision 转写。]"
            )
        return _cap_extracted_text(text, uploaded_file.name)

    if name.endswith(".docx"):
        data = _read_with_cap(uploaded_file, _MAX_CONTAINER_BYTES)
        if data is None:
            return (
                f"[DOCX 过大跳过: {uploaded_file.name} > "
                f"{_MAX_CONTAINER_BYTES // 1024 // 1024}MB]"
            )
        try:
            text = _docx_to_text(data, uploaded_file.name)
        except Exception as e:
            return f"[DOCX 解析失败: {e}]"
        if not text.strip():
            return f"[DOCX 未提取到任何文字: {uploaded_file.name}]"
        return _cap_extracted_text(text, uploaded_file.name)

    if name.endswith(_IMAGE_EXTS):
        data = _read_with_cap(uploaded_file, _MAX_IMAGE_BYTES)
        if data is None:
            return f"[图片过大跳过: {uploaded_file.name} > {_MAX_IMAGE_BYTES // 1024 // 1024}MB]"
        # ── 图片预转写(v0.26.0)──
        # 旧版做法:把图片转成 [BASE64_IMAGE:...] 占位符塞进 free_text。
        # 但 Anthropic API 不会把字符串里的 base64 当图像理解,所以
        # 下游 Crown Prince + 六部根本"看不见"图,只知道有张图存在。
        # 新版:本地起一个 Gemini Vision 调用,把图转成 OCR + 视觉描述
        # + 关键数据 的结构化文字块。Gemini 没配置 / 失败 → 降级到带
        # 警告语的占位符(transcribe_image_for_brief 内部包了所有异常)。
        from pipeline.agents.kimi_image_transcriber import (
            transcribe_image_for_brief,
        )
        # mime 简单从扩展名推断,够 Gemini API 用
        _ext_to_mime = {
            ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
            ".gif": "image/gif", ".webp": "image/webp",
        }
        _mime = next(
            (m for ext, m in _ext_to_mime.items() if name.endswith(ext)),
            "image/png",
        )
        return transcribe_image_for_brief(data, _mime, uploaded_file.name)

    return f"[不支持的文件类型: {uploaded_file.name}]"
